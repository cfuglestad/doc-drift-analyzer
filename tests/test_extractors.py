"""Tests for document text extraction and format dispatch."""

import io
from dataclasses import dataclass
from unittest.mock import MagicMock, patch

from docx import Document

from src.extractors import (
    extract_text,
    extract_text_from_docx,
    extract_text_from_pdf,
    extract_text_from_txt,
)


@dataclass
class StubUpload:
    """In-memory stand-in for a Streamlit uploaded file."""

    name: str
    content: bytes

    def read(self) -> bytes:
        return self.content


def test_extract_text_from_txt_ignores_invalid_utf8() -> None:
    assert extract_text_from_txt(b"valid\xff text") == "valid text"


@patch("src.extractors.PdfReader")
def test_extract_text_from_pdf_joins_pages_and_handles_missing_text(
    reader_class: MagicMock,
) -> None:
    first_page = MagicMock()
    first_page.extract_text.return_value = "First page"
    second_page = MagicMock()
    second_page.extract_text.return_value = None
    reader_class.return_value.pages = [first_page, second_page]

    result = extract_text_from_pdf(b"pdf bytes")

    assert result == "First page\n"
    reader_class.assert_called_once()


def test_extract_text_from_docx_joins_paragraphs() -> None:
    document = Document()
    document.add_paragraph("First paragraph")
    document.add_paragraph("Second paragraph")
    buffer = io.BytesIO()
    document.save(buffer)

    assert extract_text_from_docx(buffer.getvalue()) == "First paragraph\nSecond paragraph"


def test_extract_text_returns_empty_string_without_upload() -> None:
    assert extract_text(None) == ""


@patch("src.extractors.extract_text_from_txt", return_value="txt")
def test_extract_text_dispatches_txt_case_insensitively(txt_extractor: MagicMock) -> None:
    upload = StubUpload("DOCUMENT.TXT", b"content")

    assert extract_text(upload) == "txt"
    txt_extractor.assert_called_once_with(b"content")


@patch("src.extractors.extract_text_from_pdf", return_value="pdf")
def test_extract_text_dispatches_pdf(pdf_extractor: MagicMock) -> None:
    assert extract_text(StubUpload("document.pdf", b"content")) == "pdf"
    pdf_extractor.assert_called_once_with(b"content")


@patch("src.extractors.extract_text_from_docx", return_value="docx")
def test_extract_text_dispatches_docx(docx_extractor: MagicMock) -> None:
    assert extract_text(StubUpload("document.docx", b"content")) == "docx"
    docx_extractor.assert_called_once_with(b"content")


@patch("src.extractors.extract_text_from_txt", return_value="fallback")
def test_extract_text_uses_txt_fallback_for_unknown_extension(
    txt_extractor: MagicMock,
) -> None:
    assert extract_text(StubUpload("document.csv", b"content")) == "fallback"
    txt_extractor.assert_called_once_with(b"content")
